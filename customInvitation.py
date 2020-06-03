# customInvitaion.py - create invitaions via word document.

import docx
guest = open('guests.txt','r')
guestList = guest.readlines()
doc = docx.Document('template.docx')
for name in guestList:
    clearName = name.strip()
    paraOne = doc.add_paragraph('It would be a pleasure to have the company of')
    paraTwo = doc.add_paragraph(clearName)
    paraThree = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    paraFour = doc.add_paragraph('April 1st')
    paraFive = doc.add_paragraph("at 7 o'clock")
    paraOne.style = 'one'
    paraTwo.style = 'two'
    paraThree.style = 'three'
    paraFour.style = 'four'
    paraFive.style = 'three'
    paraFive.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.save('Invitations.docx')
